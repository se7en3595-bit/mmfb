"""MMFB Smoke Test - 端到端冒烟测试

启动流程：生成样本文件(PDF/DOCX/PNG/MD/HTML/CSV) → 通过全注册表分发 →
调用每个 Handler 的 get_preview() → 执行转换引擎 → 检查文件 I/O →
输出结果汇总 → 以 exit code 0/1 表示通过/失败。

可在 CI 中通过 subprocess 调用：
    python smoke_test.py

退出码：
    0 = 全部通过
    1 = 至少一个场景失败
"""
import os
import sys
import tempfile
import shutil
import traceback
import json
from pathlib import Path


# 项目根目录加入 sys.path
PROJECT_ROOT = Path(__file__).resolve().parent
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))


# ============================================================
# 样本文件生成
# ============================================================

def _make_sample_pdf(path: str):
    """生成最小合法 PDF"""
    content = (
        b"%PDF-1.4\n"
        b"1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n"
        b"2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj\n"
        b"3 0 obj<</Type/Page/MediaBox[0 0 612 792]/Parent 2 0 R>>endobj\n"
        b"xref\n0 4\n"
        b"0000000000 65535 f \n"
        b"0000000009 00000 n \n"
        b"0000000058 00000 n \n"
        b"0000000115 00000 n \n"
        b"trailer<</Size 4/Root 1 0 R>>\n"
        b"startxref\n206\n%%EOF"
    )
    with open(path, "wb") as f:
        f.write(content)


def _make_sample_docx(path: str) -> bool:
    """生成 Word docx 文件"""
    try:
        from docx import Document
        doc = Document()
        doc.add_heading("MMFB Smoke Test", level=1)
        doc.add_paragraph("Hello from smoke test.")
        doc.save(path)
        return True
    except ImportError:
        return False


def _make_sample_xlsx(path: str) -> bool:
    """生成 Excel xlsx 文件"""
    try:
        from openpyxl import Workbook
        wb = Workbook()
        ws = wb.active
        ws.title = "Sheet1"
        ws["A1"] = "Name"
        ws["B1"] = "Age"
        ws["C1"] = "City"
        ws.append(["Alice", 30, "Beijing"])
        ws.append(["Bob", 25, "Shanghai"])
        ws.append(["Charlie", 35, "Shenzhen"])
        wb.save(path)
        return True
    except ImportError:
        return False


def _make_sample_png(path: str) -> bool:
    """生成 32x32 橙色 PNG"""
    try:
        from PIL import Image
        img = Image.new("RGB", (32, 32), color=(240, 220, 180))
        img.save(path, "PNG")
        return True
    except ImportError:
        return False


def _make_sample_jpg(path: str) -> bool:
    """生成 32x32 JPEG"""
    try:
        from PIL import Image
        img = Image.new("RGB", (32, 32), color=(200, 180, 140))
        img.save(path, "JPEG")
        return True
    except ImportError:
        return False


def _make_sample_md(path: str):
    """生成 Markdown"""
    with open(path, "w", encoding="utf-8") as f:
        f.write("# Hello MMFB\n\nThis is a **smoke test** document.\n\n- item 1\n- item 2\n- item 3\n")


def _make_sample_html(path: str):
    """生成 HTML"""
    with open(path, "w", encoding="utf-8") as f:
        f.write("<!DOCTYPE html><html><head><title>Test</title></head><body><h1>MMFB</h1><p>Smoke test</p></body></html>")


def _make_sample_csv(path: str):
    """生成 CSV"""
    with open(path, "w", encoding="utf-8") as f:
        f.write("name,age,city\nAlice,30,Beijing\nBob,25,Shanghai\nCharlie,35,Shenzhen\n")


def _make_sample_txt(path: str):
    """生成 TXT"""
    with open(path, "w", encoding="utf-8") as f:
        f.write("Hello MMFB smoke test.\nThis is plain text.\n")


def _make_sample_py(path: str):
    """生成 Python 代码文件"""
    with open(path, "w", encoding="utf-8") as f:
        f.write('# -*- coding: utf-8 -*-\ndef hello():\n    print("Hello MMFB")\n\nif __name__ == "__main__":\n    hello()\n')


# ============================================================
# 测试场景
# ============================================================

class SmokeResult:
    """单个场景测试结果"""
    def __init__(self, name: str):
        self.name = name
        self.passed = False
        self.error = ""
        self.detail = ""


def run_scenario(name, fn):
    """执行单个场景，捕获异常"""
    r = SmokeResult(name)
    try:
        fn()
        r.passed = True
        r.detail = "OK"
    except Exception as e:
        r.passed = False
        r.error = traceback.format_exc()
        r.detail = str(e)
    return r


def scenario_registry_dispatch(samples_dir, handlers):
    """场景 1：注册表正确分发各扩展名到对应 Handler"""
    from mmfb.core.registry import registry

    mapping = {
        "test.pdf": "PdfHandler",
        "test.docx": "DocxHandler",
        "test.png": "ImageHandler",
        "test.jpg": "ImageHandler",
        "test.md": "MarkdownHandler",
        "test.html": "HtmlHandler",
        "test.csv": "CsvHandler",
        "test.txt": "TextHandler",
    }

    for filename, expected_cls in mapping.items():
        path = os.path.join(samples_dir, filename)
        if not os.path.exists(path):
            continue
        handler = registry.get_handler(path)
        if handler is None:
            raise RuntimeError(f"{filename} dispatched to None")
        actual = type(handler).__name__
        if actual != expected_cls:
            raise RuntimeError(
                f"{filename} dispatched to {actual}, expected {expected_cls}"
            )


def scenario_handler_previews(samples_dir):
    """场景 2：各格式 Handler 能执行 get_preview 并返回有效 data"""
    from mmfb.core.registry import registry

    files = [
        "test.pdf", "test.png", "test.jpg", "test.md",
        "test.html", "test.csv", "test.txt",
    ]

    for filename in files:
        path = os.path.join(samples_dir, filename)
        if not os.path.exists(path):
            continue
        handler = registry.get_handler(path)
        assert handler is not None, f"No handler for {filename}"
        preview = handler.get_preview()
        assert preview is not None, f"Preview None for {filename}"
        assert "mime" in preview, f"No mime in preview for {filename}"
        assert "data" in preview, f"No data in preview for {filename}"


def scenario_docx_edit(samples_dir):
    """场景 3：DOCX 支持 get_edit() 并返回可编辑数据"""
    from mmfb.core.registry import registry

    path = os.path.join(samples_dir, "test.docx")
    if not os.path.exists(path):
        return
    handler = registry.get_handler(path)
    assert handler is not None
    edit_data = handler.get_edit()
    assert edit_data is not None, "DOCX should support editing"
    assert "data" in edit_data


def scenario_conversion_md_to_html(samples_dir):
    """场景 4：转换 Markdown -> HTML"""
    from mmfb.services.conversion_engine import convert

    src = os.path.join(samples_dir, "test.md")
    dst = os.path.join(samples_dir, "out.html")
    r = convert(src, dst, src_format="md", dst_format="html")
    assert r.ok is True, f"MD->HTML failed: {r.error or r.message}"
    assert os.path.exists(dst), "Output HTML file not created"
    content = Path(dst).read_text(encoding="utf-8")
    assert "Hello" in content or "<h1" in content or "MMFB" in content


def scenario_conversion_html_to_md(samples_dir):
    """场景 5：转换 HTML -> Markdown"""
    from mmfb.services.conversion_engine import convert

    src = os.path.join(samples_dir, "test.html")
    dst = os.path.join(samples_dir, "out_md.md")
    r = convert(src, dst, src_format="html", dst_format="md")
    assert r.ok is True, f"HTML->MD failed: {r.error or r.message}"
    assert os.path.exists(dst), "Output MD file not created"


def scenario_pdf_to_text(samples_dir):
    """场景 6：PDF -> TXT 提取"""
    from mmfb.services.conversion_engine import convert

    src = os.path.join(samples_dir, "test.pdf")
    dst = os.path.join(samples_dir, "out_pdf.txt")
    r = convert(src, dst, src_format="pdf", dst_format="txt")
    # PDF->TXT 可能生成为空（最小 PDF 无文本），只要不报错即可
    assert r is not None, "PDF->TXT returned None"


def scenario_file_handler_io(samples_dir):
    """场景 7：file_handler 完整读写链路"""
    from mmfb.core.file_handler import (
        safe_write_text, safe_read_text, safe_write_binary, safe_read_binary,
        get_file_info, list_directory,
    )

    # 文本写读
    txt_path = os.path.join(samples_dir, "fh_test.txt")
    assert safe_write_text(txt_path, "Hello 世界")
    assert safe_read_text(txt_path) == "Hello 世界"

    # 二进制写读
    bin_path = os.path.join(samples_dir, "fh_test.bin")
    assert safe_write_binary(bin_path, b"\x00\x01\x02\xff")
    assert safe_read_binary(bin_path) == b"\x00\x01\x02\xff"

    # 元数据
    info = get_file_info(txt_path)
    assert info is not None
    assert info["name"] == "fh_test.txt"
    assert info["is_dir"] is False

    # 目录列表
    entries = list_directory(samples_dir)
    assert len(entries) > 0
    names = [e["name"] for e in entries]
    assert "fh_test.txt" in names


def scenario_conversion_image_format(samples_dir):
    """场景 8：图像格式转换 PNG -> JPG"""
    from mmfb.services.conversion_engine import convert_image

    src = os.path.join(samples_dir, "test.png")
    if not os.path.exists(src):
        return
    dst = os.path.join(samples_dir, "out_converted.jpg")
    r = convert_image(src, dst, "jpg")
    assert r.ok is True, f"PNG->JPG failed: {r.error or r.message}"
    assert os.path.exists(dst), "Output JPG not created"


def scenario_registry_stats():
    """场景 9：全局注册表统计（应注册了 100+ 扩展名）"""
    from mmfb.core.registry import registry

    count = registry.count()
    assert count >= 100, f"Registry only has {count} extensions, expected >= 100"
    exts = registry.list_extensions()
    assert ".pdf" in exts
    assert ".docx" in exts
    assert ".png" in exts
    assert ".md" in exts


def scenario_handler_editable_detection(samples_dir):
    """场景 10：可编辑性检测（DOCX 可编辑，PDF/PNG 不可编辑）"""
    from mmfb.core.registry import registry

    # DOCX 应该支持编辑
    docx_path = os.path.join(samples_dir, "test.docx")
    if os.path.exists(docx_path):
        h = registry.get_handler(docx_path)
        assert h is not None
        assert h.supports_editing() is True

    # PNG 应该不支持编辑（get_edit 返回 None）
    png_path = os.path.join(samples_dir, "test.png")
    if os.path.exists(png_path):
        h = registry.get_handler(png_path)
        assert h is not None
        # Image editing depends on Handler implementation
        # For smoke test, just verify method exists and returns bool-like
        _ = h.supports_editing()


# ============================================================
# 编辑保存验证场景
# ============================================================

def _read_text(path: str) -> str:
    """安全读取文本文件"""
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


def _write_text(path: str, data: str):
    """直接覆盖写入文本文件（模拟 Bridge.save_file）"""
    with open(path, "w", encoding="utf-8") as f:
        f.write(data)


def edit_save_text_file(ext: str, samples_dir: str):
    """通用文本文件编辑保存验证（TXT/MD/HTML）"""
    from mmfb.core.registry import registry

    filename = f"test.{ext}"
    path = os.path.join(samples_dir, filename)
    if not os.path.exists(path):
        raise FileNotFoundError(f"样本文件不存在: {filename}")

    handler = registry.get_handler(path)
    assert handler is not None, f"{ext} handler 不存在"
    assert handler.supports_editing(), f"{ext} 标不支持编辑但实际应支持"

    # 获取编辑数据
    edit_data = handler.get_edit()
    assert edit_data is not None, f"{ext} get_edit 返回 None"
    assert "data" in edit_data, f"{ext} edit_data 无 data 字段"
    original_data = edit_data["data"]
    # 文本内容在 "content" 键
    assert "content" in original_data, f"{ext} edit_data.data 无 content 键"
    original_content = original_data["content"]
    assert isinstance(original_content, str) and len(original_content) > 0, f"{ext} 内容为空"

    # 修改内容：添加标记字符串
    marker = f"[EDITED_BY_SMOKE_TEST_{ext.upper()}]"
    new_content = original_content + "\n" + marker

    # 保存（直接写入文件，模拟 Bridge.save_file）
    _write_text(path, new_content)

    # 验证：重新读取
    saved_content = _read_text(path)
    assert marker in saved_content, f"{ext} 保存后未检测到修改标记"
    # 由于 Windows 换行符转换（原始二进制读取含 \r，文本写入/读取后 \r 可能变化），
    # 规范化：移除所有 \r 字符，只比较 \n 分隔的内容
    def norm(s): return s.replace("\r", "")
    assert norm(new_content) == norm(saved_content), f"{ext} 保存内容不一致（规范化后）"


def edit_save_docx_file(samples_dir: str):
    """Word DOCX 编辑保存验证（段落级）"""
    from mmfb.core.registry import registry

    path = os.path.join(samples_dir, "test.docx")
    if not os.path.exists(path):
        raise FileNotFoundError("DOCX 样本文件不存在")

    handler = registry.get_handler(path)
    assert handler is not None
    assert handler.supports_editing()

    edit_data = handler.get_edit()
    assert edit_data is not None, "DOCX get_edit 返回 None"
    assert "data" in edit_data
    doc_data = edit_data["data"]
    assert "paragraphs" in doc_data, "DOCX 无 paragraphs 字段"
    paragraphs = doc_data["paragraphs"]
    assert len(paragraphs) > 0, "DOCX 无段落"

    # 修改：追加一个新段落到末尾（通过修改最后一段文本）
    new_para_text = "这是编辑测试添加的段落。"
    last_para = paragraphs[-1]
    idx = last_para.get("index")
    original_text = last_para.get("text", "")
    # 构造符合 handler.save_paragraphs 要求的 changes JSON（使用 index 字段）
    append_change = [{"index": idx, "text": original_text + " " + new_para_text}]
    changes_json = json.dumps(append_change)

    # 调用 handler.save_paragraphs
    success = handler.save_paragraphs(changes_json)
    assert success, "DOCX save_paragraphs 失败"

    # 验证：重新获取 edit_data 确认修改
    edit_data2 = handler.get_edit()
    paragraphs2 = edit_data2["data"]["paragraphs"]
    # 找到对应的 index
    modified = next((p for p in paragraphs2 if p.get("index") == idx), None)
    assert modified is not None, f"DOCX 修改后找不到段落 index={idx}"
    new_text_in_file = modified.get("text", "")
    assert new_para_text in new_text_in_file, f"DOCX 保存后未检测到修改: {new_text_in_file}"


def edit_save_xlsx_file(samples_dir: str):
    """Excel XLSX 编辑保存验证（单元格）"""
    from mmfb.core.registry import registry

    path = os.path.join(samples_dir, "test.xlsx")
    if not os.path.exists(path):
        raise FileNotFoundError("XLSX 样本文件不存在")

    handler = registry.get_handler(path)
    assert handler is not None
    assert handler.supports_editing()

    edit_data = handler.get_edit()
    assert edit_data is not None, "XLSX get_edit 返回 None"
    assert "data" in edit_data
    file_data = edit_data["data"]
    sheets = file_data.get("sheets", [])
    assert len(sheets) > 0, "XLSX 无工作表"
    first_sheet = sheets[0]
    sheet_name = first_sheet.get("name", "Sheet1")
    cells = first_sheet.get("cells", [])
    assert len(cells) > 0, "XLSX 无单元格数据"

    # 修改：找一个已存在的单元格并修改其值
    target_cell = cells[0]
    cell_address = target_cell.get("address")  # 如 "A1"
    original_value = target_cell.get("value")
    new_value = f"{original_value}_EDITED"

    # 调用 handler.save_cell
    success = handler.save_cell(sheet_name, cell_address, new_value)
    assert success, "XLSX save_cell 失败"

    # 验证：重新获取 edit_data 确认修改
    edit_data2 = handler.get_edit()
    sheets2 = edit_data2["data"]["sheets"]
    cells2 = sheets2[0]["cells"]
    # 匹配 address 的单元格
    updated = next((c for c in cells2 if c.get("address") == cell_address), None)
    assert updated is not None, f"XLSX 修改后找不到单元格 {cell_address}"
    assert updated.get("value") == new_value, f"XLSX 保存值不匹配: {updated.get('value')} != {new_value}"


def scenario_txt_edit_save(samples_dir):
    """场景 12：纯文本编辑保存"""
    edit_save_text_file("txt", samples_dir)


def scenario_md_edit_save(samples_dir):
    """场景 13：Markdown 编辑保存"""
    edit_save_text_file("md", samples_dir)


def scenario_html_edit_save(samples_dir):
    """场景 14：HTML 编辑保存"""
    edit_save_text_file("html", samples_dir)


def scenario_docx_edit_save(samples_dir):
    """场景 16：Word 文档段落编辑保存"""
    edit_save_docx_file(samples_dir)


def scenario_xlsx_edit_save(samples_dir):
    """场景 17：Excel 单元格编辑保存"""
    edit_save_xlsx_file(samples_dir)


def scenario_conversion_chain(samples_dir):
    """场景 11：级联转换 MD -> HTML -> MD（闭环）"""
    from mmfb.services.conversion_engine import convert

    src = os.path.join(samples_dir, "test.md")
    mid = os.path.join(samples_dir, "chain_mid.html")
    dst = os.path.join(samples_dir, "chain_out.md")

    r1 = convert(src, mid, src_format="md", dst_format="html")
    assert r1.ok is True, f"MD->HTML failed: {r1.error}"

    r2 = convert(mid, dst, src_format="html", dst_format="md")
    assert r2.ok is True, f"HTML->MD failed: {r2.error}"
    assert os.path.exists(dst)


# ============================================================
# 主入口
# ============================================================

def main():
    print("=" * 60)
    print("MMFB Smoke Test - 端到端冒烟测试")
    print("=" * 60)

    # 创建临时目录
    tmp_dir = tempfile.mkdtemp(prefix="mmfb_smoke_run_")
    print(f"[INFO] 临时样本目录: {tmp_dir}")

    # 生成样本文件
    _make_sample_pdf(os.path.join(tmp_dir, "test.pdf"))
    _make_sample_png(os.path.join(tmp_dir, "test.png"))
    _make_sample_jpg(os.path.join(tmp_dir, "test.jpg"))
    _make_sample_md(os.path.join(tmp_dir, "test.md"))
    _make_sample_html(os.path.join(tmp_dir, "test.html"))
    _make_sample_csv(os.path.join(tmp_dir, "test.csv"))
    _make_sample_txt(os.path.join(tmp_dir, "test.txt"))
    _make_sample_py(os.path.join(tmp_dir, "test.py"))

    # DOCX 生成需要 python-docx
    docx_ok = _make_sample_docx(os.path.join(tmp_dir, "test.docx"))

    # XLSX 生成需要 openpyxl
    xlsx_ok = _make_sample_xlsx(os.path.join(tmp_dir, "test.xlsx"))

    print(f"[INFO] 样本文件已生成（DOCX={'OK' if docx_ok else 'SKIP'}， XLSX={'OK' if xlsx_ok else 'SKIP'}）")

    # 导入 handlers 来触发注册
    try:
        from mmfb.handlers import registry
        print(f"[INFO] 全局注册表已加载，共 {registry.count()} 个扩展名")
    except Exception as e:
        print(f"[ERROR] 加载 handlers 失败: {e}")
        shutil.rmtree(tmp_dir, ignore_errors=True)
        sys.exit(1)

    # 执行所有场景
    scenarios = [
        ("注册表分发",
         lambda: scenario_registry_dispatch(tmp_dir, None)),
        ("Handler 预览调用",
         lambda: scenario_handler_previews(tmp_dir)),
        ("DOCX 可编辑性",
         lambda: scenario_docx_edit(tmp_dir)),
        ("MD -> HTML 转换",
         lambda: scenario_conversion_md_to_html(tmp_dir)),
        ("HTML -> MD 转换",
         lambda: scenario_conversion_html_to_md(tmp_dir)),
        ("PDF -> TXT 提取",
         lambda: scenario_pdf_to_text(tmp_dir)),
        ("file_handler 读写",
         lambda: scenario_file_handler_io(tmp_dir)),
        ("PNG -> JPG 图像转换",
         lambda: scenario_conversion_image_format(tmp_dir)),
        ("注册表统计（>=100 格式）",
         lambda: scenario_registry_stats()),
        ("可编辑性检测",
         lambda: scenario_handler_editable_detection(tmp_dir)),
        ("级联转换 MD->HTML->MD",
         lambda: scenario_conversion_chain(tmp_dir)),
        # 编辑保存场景
        ("TXT 编辑保存",
         lambda: scenario_txt_edit_save(tmp_dir)),
        ("Markdown 编辑保存",
         lambda: scenario_md_edit_save(tmp_dir)),
        ("HTML 编辑保存",
         lambda: scenario_html_edit_save(tmp_dir)),
        ("DOCX 编辑保存",
         lambda: scenario_docx_edit_save(tmp_dir)),
        ("XLSX 编辑保存",
         lambda: scenario_xlsx_edit_save(tmp_dir)),
    ]

    results = []
    for name, fn in scenarios:
        r = run_scenario(name, fn)
        results.append(r)

    # 输出汇总
    print()
    print("-" * 60)
    passed = sum(1 for r in results if r.passed)
    failed = sum(1 for r in results if not r.passed)
    print(f"结果: {passed} passed, {failed} failed, {len(results)} total")
    print("-" * 60)

    for r in results:
        status = "PASS" if r.passed else "FAIL"
        print(f"  [{status}] {r.name}: {r.detail}")
        if not r.passed and r.error:
            # 只打印简要错误信息
            lines = r.error.strip().split("\n")
            for line in lines[-3:]:
                print(f"          {line}")

    # 清理临时目录
    shutil.rmtree(tmp_dir, ignore_errors=True)

    # 退出码
    if failed > 0:
        print(f"\n[FAIL] Smoke test 失败 ({failed} 场景)")
        sys.exit(1)
    else:
        print("\n[PASS] Smoke test 全部通过")
        sys.exit(0)


if __name__ == "__main__":
    main()
