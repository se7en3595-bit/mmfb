"""ConversionEngine - 统一格式转换引擎

支持的转换链路：
  文档类 (Task #14):
    MD->HTML        markdown 库
    HTML->MD        html2text 库
    MD->DOCX        pandoc (优先) 或 python-docx 回退
    DOCX->HTML      mammoth
    DOCX->MD        mammoth
    HTML->PDF       PySide6 QPrinter
    PDF->TXT        pypdf 文本提取
    PDF->MD         pypdf 文本提取 + Markdown 组装

  表格类 (Task #13, 本次新增):
    XLSX->CSV       openpyxl + csv
    CSV->XLSX       pandas + openpyxl
    XLSX->TSV       openpyxl + csv
    CSV->CSV        pandas (重格式化/编码转换)

  图像类 (Task #13, 本次新增):
    PNG<->JPG<->WebP<->BMP<->TIFF<->ICO   (Pillow)
    支持：位深自动转换、质量参数、EXIF 保留

  Base64 编解码 (Task #13, 本次新增):
    <任意文件>->TXT (Base64 编码)
    TXT-><原格式> (Base64 解码)

设计：
  - 所有方法都是纯函数，返回 ConversionResult
  - convert() 统一派发，新增 progress_cb 可选进度回调
  - get_supported_conversions() 返回支持的转换列表
"""
import base64
import csv
import os
import shutil
import subprocess
import sys
from dataclasses import dataclass, field
from typing import Callable, Optional


@dataclass
class ConversionResult:
    """转换结果"""
    ok: bool
    output_path: str = ""
    message: str = ""
    error: str = ""
    metadata: dict = field(default_factory=dict)


# ============================================================
# 内部辅助
# ============================================================

def _find_exe(name: str) -> Optional[str]:
    return shutil.which(name)


def _run(args, timeout=60):
    """运行子进程，Windows 下隐藏控制台窗口"""
    startupinfo = None
    if sys.platform == "win32":
        startupinfo = subprocess.STARTUPINFO()
        startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW
    return subprocess.run(
        args,
        capture_output=True,
        text=True,
        timeout=timeout,
        startupinfo=startupinfo,
        encoding="utf-8",
        errors="replace",
    )


def _ensure_parent_dir(path: str):
    d = os.path.dirname(path)
    if d and not os.path.isdir(d):
        os.makedirs(d, exist_ok=True)


def _unique_path(path: str) -> str:
    """返回不冲突的文件路径（序号递增）"""
    if not os.path.exists(path):
        return path
    base, ext = os.path.splitext(path)
    idx = 1
    while os.path.exists(f"{base}_{idx}{ext}"):
        idx += 1
    return f"{base}_{idx}{ext}"


# ============================================================
# MD -> HTML
# ============================================================

def md_to_html(source_path: str, output_path: str) -> ConversionResult:
    """Markdown -> HTML"""
    try:
        import markdown
    except ImportError:
        return ConversionResult(ok=False, error="markdown library not installed")

    try:
        with open(source_path, "r", encoding="utf-8", errors="replace") as f:
            md_text = f.read()
    except Exception as e:
        return ConversionResult(ok=False, error=f"read source failed: {e}")

    try:
        body_html = markdown.markdown(
            md_text,
            extensions=["fenced_code", "tables", "toc", "nl2br"],
        )
    except Exception as e:
        return ConversionResult(ok=False, error=f"markdown conversion failed: {e}")

    full_html = (
        "<!DOCTYPE html>\n"
        "<html lang=\"zh-CN\">\n"
        "<head>\n"
        "<meta charset=\"utf-8\">\n"
        "<style>\n"
        "body{font-family:system-ui,sans-serif;max-width:800px;margin:2em auto;"
        "line-height:1.7;color:#333;}\n"
        "pre{background:#f6f8fa;padding:1em;border-radius:6px;overflow-x:auto;}\n"
        "code{background:#f6f8fa;padding:0.15em 0.4em;border-radius:3px;}\n"
        "table{border-collapse:collapse;width:100%;}\n"
        "th,td{border:1px solid #ddd;padding:0.5em;text-align:left;}\n"
        "img{max-width:100%;}\n"
        "blockquote{border-left:4px solid #ddd;margin-left:0;padding-left:1em;color:#666;}\n"
        "</style>\n"
        "</head>\n"
        "<body>\n"
        + body_html +
        "\n</body>\n</html>"
    )

    _ensure_parent_dir(output_path)
    try:
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(full_html)
    except Exception as e:
        return ConversionResult(ok=False, error=f"write output failed: {e}")

    return ConversionResult(
        ok=True, output_path=output_path,
        message=f"MD->HTML success: {os.path.basename(output_path)}",
    )


# ============================================================
# HTML -> MD
# ============================================================

def html_to_md(source_path: str, output_path: str) -> ConversionResult:
    """HTML -> Markdown"""
    try:
        import html2text
    except ImportError:
        return ConversionResult(ok=False, error="html2text library not installed")

    try:
        with open(source_path, "r", encoding="utf-8", errors="replace") as f:
            html_text = f.read()
    except Exception as e:
        return ConversionResult(ok=False, error=f"read source failed: {e}")

    try:
        converter = html2text.HTML2Text()
        converter.body_width = 0
        converter.ignore_links = False
        converter.ignore_images = False
        converter.single_line_break = False
        md_text = converter.handle(html_text)
    except Exception as e:
        return ConversionResult(ok=False, error=f"html2text failed: {e}")

    _ensure_parent_dir(output_path)
    try:
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(md_text)
    except Exception as e:
        return ConversionResult(ok=False, error=f"write output failed: {e}")

    return ConversionResult(
        ok=True, output_path=output_path,
        message=f"HTML->MD success: {os.path.basename(output_path)}",
    )


# ============================================================
# MD -> DOCX
# ============================================================

def md_to_docx(source_path: str, output_path: str) -> ConversionResult:
    """Markdown -> Word .docx（pandoc 优先）"""
    pandoc_path = _find_exe("pandoc")
    if pandoc_path:
        return _md_to_docx_pandoc(source_path, output_path, pandoc_path)
    return _md_to_docx_python(source_path, output_path)


def _md_to_docx_pandoc(source_path, output_path, pandoc_path) -> ConversionResult:
    _ensure_parent_dir(output_path)
    try:
        r = _run([pandoc_path, source_path, "-o", output_path, "--from=markdown", "--to=docx"])
        if r.returncode != 0:
            return ConversionResult(
                ok=False, error=f"pandoc failed (code {r.returncode}): {r.stderr[:300]}"
            )
        if not os.path.isfile(output_path):
            return ConversionResult(ok=False, error="pandoc produced no output")
        return ConversionResult(
            ok=True, output_path=output_path,
            message=f"MD->DOCX (pandoc) success: {os.path.basename(output_path)}",
        )
    except subprocess.TimeoutExpired:
        return ConversionResult(ok=False, error="pandoc timed out")
    except Exception as e:
        return ConversionResult(ok=False, error=f"pandoc error: {e}")


def _md_to_docx_python(source_path, output_path) -> ConversionResult:
    """python-docx 回退实现"""
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        return ConversionResult(ok=False, error="python-docx library not installed")

    try:
        with open(source_path, "r", encoding="utf-8", errors="replace") as f:
            md_text = f.read()
    except Exception as e:
        return ConversionResult(ok=False, error=f"read source failed: {e}")

    doc = Document()
    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            title_text = stripped[level:].strip()
            if title_text and 1 <= level <= 6:
                h = doc.add_heading(level=level)
                _add_inline_runs(h, title_text)
                i += 1
                continue

        if stripped.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1
            p = doc.add_paragraph()
            run = p.add_run("\n".join(code_lines))
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            continue

        if stripped.startswith("- ") or stripped.startswith("* "):
            while i < len(lines) and lines[i].strip().startswith(("- ", "* ")):
                item_text = lines[i].strip()[2:]
                _add_list_item(doc, item_text)
                i += 1
            continue

        if stripped and stripped[0].isdigit() and ". " in stripped[:4]:
            while (i < len(lines) and lines[i].strip()
                   and lines[i].strip()[0].isdigit()
                   and ". " in lines[i].strip()[:4]):
                dot_pos = lines[i].strip().index(". ")
                item_text = lines[i].strip()[dot_pos + 2:]
                p = doc.add_paragraph(style="List Number")
                _add_inline_runs(p, item_text)
                i += 1
            continue

        if not stripped:
            i += 1
            continue

        p = doc.add_paragraph()
        _add_inline_runs(p, stripped)
        i += 1

    _ensure_parent_dir(output_path)
    try:
        doc.save(output_path)
    except Exception as e:
        return ConversionResult(ok=False, error=f"write output failed: {e}")

    return ConversionResult(
        ok=True, output_path=output_path,
        message=f"MD->DOCX (python-docx) success: {os.path.basename(output_path)}",
    )


def _add_list_item(doc, text: str):
    p = doc.add_paragraph(style="List Bullet")
    _add_inline_runs(p, text)


def _add_inline_runs(paragraph, text: str):
    """解析简单 markdown 行内语法"""
    import re
    pattern = re.compile(r'(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)')
    parts = []
    last_end = 0
    for m in pattern.finditer(text):
        if m.start() > last_end:
            parts.append(("text", text[last_end:m.start()]))
        token = m.group(0)
        if token.startswith("**"):
            parts.append(("bold", token[2:-2]))
        elif token.startswith("`"):
            parts.append(("code", token[1:-1]))
        elif token.startswith("*"):
            parts.append(("italic", token[1:-1]))
        else:
            parts.append(("text", token))
        last_end = m.end()
    if last_end < len(text):
        parts.append(("text", text[last_end:]))

    if not parts and text:
        parts.append(("text", text))

    from docx.shared import Pt
    for part in parts:
        if part[0] == "text":
            paragraph.add_run(part[1])
        elif part[0] == "bold":
            run = paragraph.add_run(part[1])
            run.bold = True
        elif part[0] == "italic":
            run = paragraph.add_run(part[1])
            run.italic = True
        elif part[0] == "code":
            run = paragraph.add_run(part[1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)


# ============================================================
# DOCX -> HTML / MD
# ============================================================

def docx_to_html(source_path: str, output_path: str) -> ConversionResult:
    """DOCX -> HTML"""
    import mammoth
    if not os.path.isfile(source_path):
        return ConversionResult(ok=False, error=f"source not found: {source_path}")
    try:
        with open(source_path, "rb") as f:
            result = mammoth.convert_to_html(f)
            body_html = result.value
    except Exception as e:
        return ConversionResult(ok=False, error=f"mammoth failed: {e}")

    warnings = getattr(result, "messages", [])
    warn_messages = [getattr(m, "message", str(m)) for m in warnings] if warnings else []

    full_html = (
        "<!DOCTYPE html>\n<html lang=\"zh-CN\">\n<head>\n<meta charset=\"utf-8\">\n"
        "<style>\n"
        "body{font-family:system-ui,sans-serif;max-width:800px;margin:2em auto;"
        "line-height:1.7;color:#333;}\n"
        "table{border-collapse:collapse;width:100%;}\n"
        "th,td{border:1px solid #ddd;padding:0.5em;text-align:left;}\n"
        "img{max-width:100%;}\n"
        "</style>\n</head>\n<body>\n"
        + body_html + "\n</body>\n</html>"
    )
    _ensure_parent_dir(output_path)
    try:
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(full_html)
    except Exception as e:
        return ConversionResult(ok=False, error=f"write output failed: {e}")
    return ConversionResult(
        ok=True, output_path=output_path,
        message=f"DOCX->HTML success: {os.path.basename(output_path)}",
        metadata={"warnings": warn_messages},
    )


def docx_to_md(source_path: str, output_path: str) -> ConversionResult:
    """DOCX -> Markdown"""
    import mammoth
    if not os.path.isfile(source_path):
        return ConversionResult(ok=False, error=f"source not found: {source_path}")
    try:
        with open(source_path, "rb") as f:
            result = mammoth.convert_to_markdown(f)
            md_text = result.value
    except Exception as e:
        return ConversionResult(ok=False, error=f"mammoth failed: {e}")
    _ensure_parent_dir(output_path)
    try:
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(md_text)
    except Exception as e:
        return ConversionResult(ok=False, error=f"write output failed: {e}")
    return ConversionResult(
        ok=True, output_path=output_path,
        message=f"DOCX->MD success: {os.path.basename(output_path)}",
    )


# ============================================================
# HTML -> PDF (QPrinter)
# ============================================================

def html_to_pdf(source_path: str, output_path: str) -> ConversionResult:
    """HTML -> PDF via wkhtmltopdf or weasyprint（纯子进程，无 QWebEngineView）

    不再使用 QWebEngineView + QEventLoop，避免阻塞 UI 线程和崩溃问题。
    """
    _ensure_parent_dir(output_path)

    # 1. 优先使用 wkhtmltopdf（命令行工具，轻量稳定）
    wkhtml_path = shutil.which("wkhtmltopdf")
    if wkhtml_path:
        try:
            startupinfo = None
            if sys.platform == "win32":
                startupinfo = subprocess.STARTUPINFO()
                startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW
            proc = subprocess.run(
                [wkhtml_path, "--quiet", "--enable-local-file-access",
                 os.path.abspath(source_path), os.path.abspath(output_path)],
                capture_output=True, text=True, timeout=60,
                startupinfo=startupinfo,
            )
            if proc.returncode == 0 and os.path.isfile(output_path) and os.path.getsize(output_path) > 0:
                return ConversionResult(
                    ok=True, output_path=output_path,
                    message=f"HTML->PDF success: {os.path.basename(output_path)}",
                )
            return ConversionResult(ok=False, error=f"wkhtmltopdf failed: {proc.stderr[:200]}")
        except subprocess.TimeoutExpired:
            return ConversionResult(ok=False, error="wkhtmltopdf timeout (60s)")
        except Exception as e:
            return ConversionResult(ok=False, error=f"wkhtmltopdf error: {e}")

    # 2. 回退到 weasyprint
    try:
        import weasyprint
        html_doc = weasyprint.HTML(filename=os.path.abspath(source_path))
        html_doc.write_pdf(os.path.abspath(output_path))
        if os.path.isfile(output_path) and os.path.getsize(output_path) > 0:
            return ConversionResult(
                ok=True, output_path=output_path,
                message=f"HTML->PDF success: {os.path.basename(output_path)}",
            )
        return ConversionResult(ok=False, error="weasyprint output is empty")
    except ImportError:
        pass
    except Exception as e:
        return ConversionResult(ok=False, error=f"weasyprint error: {e}")

    # 3. 都不支持
    return ConversionResult(
        ok=False,
        error="HTML->PDF 需要安装 wkhtmltopdf 或 weasyprint（pip install weasyprint）",
    )


# ============================================================
# PDF -> TXT / MD / PNG
# ============================================================

def pdf_to_text(source_path: str, output_path: str,
                progress_cb: Optional[Callable[[int, int], None]] = None) -> ConversionResult:
    """PDF -> TXT 纯文本提取（支持进度回调）"""
    from pypdf import PdfReader
    if not os.path.isfile(source_path):
        return ConversionResult(ok=False, error=f"source not found: {source_path}")
    try:
        reader = PdfReader(source_path)
    except Exception as e:
        return ConversionResult(ok=False, error=f"pypdf failed to open: {e}")

    pages_text = []
    total = len(reader.pages)
    for i, page in enumerate(reader.pages, start=1):
        try:
            t = page.extract_text()
            if t:
                pages_text.append(t)
        except Exception:
            continue
        if progress_cb and total and i % 10 == 0:
            progress_cb(i, total)

    if progress_cb and total:
        progress_cb(total, total)

    text = "\n\n".join(pages_text)
    _ensure_parent_dir(output_path)
    try:
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(text)
    except Exception as e:
        return ConversionResult(ok=False, error=f"write output failed: {e}")

    return ConversionResult(
        ok=True, output_path=output_path,
        message=f"PDF->TXT success: {total} pages, {len(text)} chars",
        metadata={"page_count": total, "char_count": len(text)},
    )


def pdf_to_md(source_path: str, output_path: str,
              progress_cb: Optional[Callable[[int, int], None]] = None) -> ConversionResult:
    """PDF -> Markdown（按页分组，结构化输出，支持进度回调）"""
    from pypdf import PdfReader
    if not os.path.isfile(source_path):
        return ConversionResult(ok=False, error=f"source not found: {source_path}")
    try:
        reader = PdfReader(source_path)
    except Exception as e:
        return ConversionResult(ok=False, error=f"pypdf failed to open: {e}")

    full_md_parts = []
    total = len(reader.pages)
    for i, page in enumerate(reader.pages, start=1):
        try:
            t = page.extract_text() or ""
        except Exception:
            t = ""
        full_md_parts.append(f"## Page {i}\n\n{t}\n")
        if progress_cb and total and i % 10 == 0:
            progress_cb(i, total)

    if progress_cb and total:
        progress_cb(total, total)

    md_text = "\n".join(full_md_parts).strip()
    _ensure_parent_dir(output_path)
    try:
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(md_text)
    except Exception as e:
        return ConversionResult(ok=False, error=f"write output failed: {e}")

    return ConversionResult(
        ok=True, output_path=output_path,
        message=f"PDF->MD success: {total} pages",
        metadata={"page_count": total},
    )


def pdf_to_png(source_path: str, output_path: str,
               dpi: int = 150,
               progress_cb: Optional[Callable[[int, int], None]] = None) -> ConversionResult:
    """PDF 单页 -> PNG 图像（PyMuPDF 渲染）

    参数：
        dpi: 渲染分辨率，默认 150（兼顾速度与清晰度）
    """
    try:
        import fitz  # PyMuPDF
    except ImportError:
        return ConversionResult(ok=False, error="PyMuPDF library not installed")

    if not os.path.isfile(source_path):
        return ConversionResult(ok=False, error=f"source not found: {source_path}")

    try:
        doc = fitz.open(source_path)
    except Exception as e:
        return ConversionResult(ok=False, error=f"PyMuPDF failed to open: {e}")

    if len(doc) == 0:
        doc.close()
        return ConversionResult(ok=False, error="PDF has no pages")

    # 只渲染第一页
    page = doc[0]
    zoom = dpi / 72.0
    mat = fitz.Matrix(zoom, zoom)
    try:
        pix = page.get_pixmap(matrix=mat)
    except Exception as e:
        doc.close()
        return ConversionResult(ok=False, error=f"render failed: {e}")

    page_count = len(doc)

    if progress_cb:
        progress_cb(1, page_count)

    _ensure_parent_dir(output_path)
    try:
        pix.save(output_path)
    except Exception as e:
        doc.close()
        return ConversionResult(ok=False, error=f"save failed: {e}")

    doc.close()
    if progress_cb:
        progress_cb(1, page_count)

    return ConversionResult(
        ok=True, output_path=output_path,
        message=f"PDF->PNG success: page 1/{page_count}, {pix.width}x{pix.height}",
        metadata={"page_count": page_count, "width": pix.width, "height": pix.height},
    )


def pdf_to_png_folder(source_path: str, output_folder: str,
                       dpi: int = 150,
                       progress_cb: Optional[Callable[[int, int], None]] = None) -> ConversionResult:
    """PDF 多页 -> PNG 文件夹（每页一个 PNG，支持进度回调）

    参数：
        dpi: 渲染分辨率，默认 150
    返回：
        output_folder 为创建的目录路径（每个文件 page_N.png）
    """
    try:
        import fitz
    except ImportError:
        return ConversionResult(ok=False, error="PyMuPDF library not installed")

    if not os.path.isfile(source_path):
        return ConversionResult(ok=False, error=f"source not found: {source_path}")

    try:
        doc = fitz.open(source_path)
    except Exception as e:
        return ConversionResult(ok=False, error=f"PyMuPDF failed to open: {e}")

    page_count = len(doc)
    if page_count == 0:
        doc.close()
        return ConversionResult(ok=False, error="PDF has no pages")

    # 创建输出目录（使用源文件名作为文件夹名）
    base_name = os.path.splitext(os.path.basename(source_path))[0]
    if not output_folder:
        output_folder = os.path.join(os.path.dirname(source_path), base_name + "_pages")
    if not os.path.isdir(output_folder):
        os.makedirs(output_folder, exist_ok=True)

    rendered = 0
    zoom = dpi / 72.0
    mat = fitz.Matrix(zoom, zoom)

    for i in range(page_count):
        page = doc[i]
        try:
            pix = page.get_pixmap(matrix=mat)
        except Exception:
            continue

        out_name = f"page_{i + 1:03d}.png"
        out_path = os.path.join(output_folder, out_name)
        try:
            pix.save(out_path)
            rendered += 1
        except Exception:
            continue

        if progress_cb and page_count and (i + 1) % 10 == 0:
            progress_cb(i + 1, page_count)

    if progress_cb and page_count:
        progress_cb(page_count, page_count)

    doc.close()

    return ConversionResult(
        ok=True, output_path=output_folder,
        message=f"PDF->PNG folder success: {rendered}/{page_count} pages rendered to {output_folder}",
        metadata={"page_count": page_count, "rendered": rendered, "dpi": dpi},
    )


# ============================================================
# XLSX -> CSV / TSV (新增)
# ============================================================

def xlsx_to_csv(source_path: str, output_path: str,
                progress_cb: Optional[Callable[[int, int], None]] = None) -> ConversionResult:
    """XLSX -> CSV（openpyxl + csv）"""
    try:
        from openpyxl import load_workbook
    except ImportError:
        return ConversionResult(ok=False, error="openpyxl library not installed")

    if not os.path.isfile(source_path):
        return ConversionResult(ok=False, error=f"source not found: {source_path}")
    try:
        wb = load_workbook(source_path, read_only=True, data_only=True)
    except Exception as e:
        return ConversionResult(ok=False, error=f"open failed: {e}")

    ws = wb.active
    if ws is None:
        # 尝试第一个 sheet
        if wb.sheetnames:
            ws = wb[wb.sheetnames[0]]
        else:
            return ConversionResult(ok=False, error="no sheet in workbook")

    _ensure_parent_dir(output_path)
    try:
        # 计算总行数用于进度
        total = ws.max_row or 0
        with open(output_path, "w", encoding="utf-8-sig", newline="") as f:
            writer = csv.writer(f)
            for i, row in enumerate(ws.iter_rows(values_only=True), start=1):
                writer.writerow(["" if v is None else v for v in row])
                if progress_cb and total and i % 500 == 0:
                    progress_cb(i, total)
        if progress_cb and total:
            progress_cb(total, total)
    except Exception as e:
        return ConversionResult(ok=False, error=f"write failed: {e}")
    finally:
        wb.close()

    return ConversionResult(
        ok=True, output_path=output_path,
        message=f"XLSX->CSV success: {os.path.basename(output_path)}",
        metadata={"rows": ws.max_row or 0, "sheet": ws.title},
    )


def xlsx_to_tsv(source_path: str, output_path: str,
                progress_cb: Optional[Callable[[int, int], None]] = None) -> ConversionResult:
    """XLSX -> TSV"""
    try:
        from openpyxl import load_workbook
    except ImportError:
        return ConversionResult(ok=False, error="openpyxl library not installed")

    if not os.path.isfile(source_path):
        return ConversionResult(ok=False, error=f"source not found: {source_path}")
    try:
        wb = load_workbook(source_path, read_only=True, data_only=True)
    except Exception as e:
        return ConversionResult(ok=False, error=f"open failed: {e}")

    ws = wb.active or (wb[wb.sheetnames[0]] if wb.sheetnames else None)
    if ws is None:
        return ConversionResult(ok=False, error="no sheet in workbook")

    _ensure_parent_dir(output_path)
    try:
        total = ws.max_row or 0
        with open(output_path, "w", encoding="utf-8", newline="") as f:
            writer = csv.writer(f, delimiter="\t")
            for i, row in enumerate(ws.iter_rows(values_only=True), start=1):
                writer.writerow(["" if v is None else v for v in row])
                if progress_cb and total and i % 500 == 0:
                    progress_cb(i, total)
        if progress_cb and total:
            progress_cb(total, total)
    except Exception as e:
        return ConversionResult(ok=False, error=f"write failed: {e}")
    finally:
        wb.close()

    return ConversionResult(
        ok=True, output_path=output_path,
        message=f"XLSX->TSV success: {os.path.basename(output_path)}",
        metadata={"rows": ws.max_row or 0, "sheet": ws.title},
    )


# ============================================================
# CSV -> XLSX / CSV -> CSV (pandas 重格式化 / 编码转换)
# ============================================================

def csv_to_xlsx(source_path: str, output_path: str,
                progress_cb: Optional[Callable[[int, int], None]] = None) -> ConversionResult:
    """CSV/TSV -> XLSX（pandas + openpyxl）"""
    try:
        import pandas as pd
    except ImportError:
        return ConversionResult(ok=False, error="pandas library not installed")

    if not os.path.isfile(source_path):
        return ConversionResult(ok=False, error=f"source not found: {source_path}")

    # 探测编码和分隔符
    enc = _detect_encoding_simple(source_path)
    delim = _detect_delimiter_simple(source_path, enc)

    try:
        df = pd.read_csv(source_path, encoding=enc, sep=delim, on_bad_lines="warn")
    except Exception as e:
        return ConversionResult(ok=False, error=f"parse failed: {e}")

    if progress_cb:
        progress_cb(50, 100)

    _ensure_parent_dir(output_path)
    try:
        df.to_excel(output_path, index=False, engine="openpyxl")
    except Exception as e:
        return ConversionResult(ok=False, error=f"write failed: {e}")

    if progress_cb:
        progress_cb(100, 100)

    return ConversionResult(
        ok=True, output_path=output_path,
        message=f"CSV->XLSX success: {os.path.basename(output_path)}",
        metadata={"rows": len(df), "cols": len(df.columns)},
    )


def csv_to_csv(source_path: str, output_path: str,
               progress_cb: Optional[Callable[[int, int], None]] = None) -> ConversionResult:
    """CSV 重格式化（UTF-8-BOM 编码统一+逗号分隔符）"""
    try:
        import pandas as pd
    except ImportError:
        return ConversionResult(ok=False, error="pandas library not installed")

    if not os.path.isfile(source_path):
        return ConversionResult(ok=False, error=f"source not found: {source_path}")

    enc = _detect_encoding_simple(source_path)
    delim = _detect_delimiter_simple(source_path, enc)

    try:
        df = pd.read_csv(source_path, encoding=enc, sep=delim, on_bad_lines="warn")
    except Exception as e:
        return ConversionResult(ok=False, error=f"parse failed: {e}")

    _ensure_parent_dir(output_path)
    try:
        df.to_csv(output_path, index=False, encoding="utf-8-sig")
    except Exception as e:
        return ConversionResult(ok=False, error=f"write failed: {e}")

    return ConversionResult(
        ok=True, output_path=output_path,
        message=f"CSV->CSV (reformatted) success: {os.path.basename(output_path)}",
        metadata={"rows": len(df), "cols": len(df.columns)},
    )


# ============================================================
# 图像互转 (PNG<->JPG<->WebP<->BMP<->TIFF<->ICO)
# ============================================================

# 支持的图像格式扩展名 -> Pillow 保存格式名
IMAGE_FORMAT_MAP = {
    "png": "PNG",
    "jpg": "JPEG",
    "jpeg": "JPEG",
    "webp": "WEBP",
    "bmp": "BMP",
    "tiff": "TIFF",
    "tif": "TIFF",
    "gif": "GIF",
    "ico": "ICO",
}

# 需要 RGBA->RGB 处理的格式
JPEG_INCOMPAT_MODES = {"RGBA", "P", "LA", "PA"}
JPEG_INCOMPAT_EXT = {"jpg", "jpeg", "bmp"}

# 格式对应的保存参数
SAVE_PARAMS = {
    "JPEG": {"quality": 92, "optimize": True},
    "WEBP": {"quality": 85, "method": 4},
    "PNG": {"optimize": True},
    "TIFF": {},
    "BMP": {},
    "GIF": {},
    "ICO": {},
}


def convert_image(source_path: str, output_path: str, dst_format: str = "",
                  quality: int = 0,
                  progress_cb: Optional[Callable[[int, int], None]] = None) -> ConversionResult:
    """图像格式互转（Pillow）

    参数：
        quality: JPEG/WebP 质量（1-100），0 使用默认值
    """
    try:
        from PIL import Image
    except ImportError:
        return ConversionResult(ok=False, error="Pillow library not installed")

    if not os.path.isfile(source_path):
        return ConversionResult(ok=False, error=f"source not found: {source_path}")

    if not dst_format:
        dst_format = os.path.splitext(output_path)[1].lstrip(".").lower()

    pillow_fmt = IMAGE_FORMAT_MAP.get(dst_format)
    if pillow_fmt is None:
        return ConversionResult(
            ok=False,
            error=f"unsupported target format: {dst_format}",
            metadata={"supported": list(IMAGE_FORMAT_MAP.keys())},
        )

    try:
        img = Image.open(source_path)
    except Exception as e:
        return ConversionResult(ok=False, error=f"open failed: {e}")

    if progress_cb:
        progress_cb(30, 100)

    # 处理模式不兼容
    save_kwargs = dict(SAVE_PARAMS.get(pillow_fmt, {}))
    if quality >= 1 and quality <= 100 and pillow_fmt in ("JPEG", "WEBP"):
        save_kwargs["quality"] = quality

    # RGBA/P 模式目标 JPEG/BMP 需转 RGB
    if pillow_fmt in ("JPEG", "BMP") and img.mode in JPEG_INCOMPAT_MODES:
        if img.mode == "RGBA":
            bg = Image.new("RGB", img.size, (255, 255, 255))
            bg.paste(img, mask=img.split()[3])
            img = bg
        else:
            img = img.convert("RGB")
    elif pillow_fmt == "ICO":
        # ICO 推荐 RGBA
        if img.mode != "RGBA":
            img = img.convert("RGBA")
        # 默认使用 256 尺寸
        if img.size[0] > 256 or img.size[1] > 256:
            img = img.resize((256, 256), Image.LANCZOS)

    if progress_cb:
        progress_cb(60, 100)

    _ensure_parent_dir(output_path)
    try:
        # GIF 动画仅保存第一帧
        img.save(output_path, format=pillow_fmt, **save_kwargs)
    except Exception as e:
        return ConversionResult(ok=False, error=f"save failed: {e}")

    if progress_cb:
        progress_cb(100, 100)

    return ConversionResult(
        ok=True, output_path=output_path,
        message=f"image->{pillow_fmt} success: {os.path.basename(output_path)}",
        metadata={
            "src_format": img.format,
            "dst_format": pillow_fmt,
            "width": img.width,
            "height": img.height,
        },
    )


# ============================================================
# Base64 编解码
# ============================================================

def file_to_base64(source_path: str, output_path: str,
                   progress_cb: Optional[Callable[[int, int], None]] = None) -> ConversionResult:
    """任意文件 -> Base64 编码 TXT"""
    if not os.path.isfile(source_path):
        return ConversionResult(ok=False, error=f"source not found: {source_path}")

    file_size = os.path.getsize(source_path)
    if file_size > 50 * 1024 * 1024:
        return ConversionResult(ok=False, error="file too large (>50MB)")

    _ensure_parent_dir(output_path)
    try:
        with open(source_path, "rb") as f:
            data = f.read()
        encoded = base64.b64encode(data).decode("ascii")

        # 末尾附加元信息注释
        src_name = os.path.basename(source_path)
        src_ext = os.path.splitext(source_path)[1]
        payload = (
            f"# MMFB Base64 Encoded\n"
            f"# source: {src_name}\n"
            f"# original_ext: {src_ext}\n"
            f"# encoding: base64\n"
            f"# size: {file_size}\n\n"
            f"{encoded}\n"
        )
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(payload)

        if progress_cb:
            progress_cb(100, 100)
    except Exception as e:
        return ConversionResult(ok=False, error=f"encode failed: {e}")

    return ConversionResult(
        ok=True, output_path=output_path,
        message=f"file->base64 success: {len(encoded)} chars",
        metadata={"size": file_size, "encoded_chars": len(encoded)},
    )


def base64_to_file(source_path: str, output_path: str,
                   progress_cb: Optional[Callable[[int, int], None]] = None) -> ConversionResult:
    """Base64 TXT -> 解码为原格式文件"""
    if not os.path.isfile(source_path):
        return ConversionResult(ok=False, error=f"source not found: {source_path}")

    try:
        with open(source_path, "r",encoding="utf-8", errors="replace") as f:
            content = f.read()
    except Exception as e:
        return ConversionResult(ok=False, error=f"read failed: {e}")

    # 解析头部元信息
    src_ext = ".bin"
    for line in content.splitlines()[:10]:
        if line.startswith("# original_ext:"):
            src_ext = line.split(":", 1)[1].strip()
            if not src_ext.startswith("."):
                src_ext = "." + src_ext
            break

    # 提取 Base64 数据（跳过注释行和空行，拼接所有纯 base64 行）
    b64_lines = []
    for line in content.splitlines():
        s = line.strip()
        if not s or s.startswith("#"):
            continue
        b64_lines.append(s)
    b64_text = "".join(b64_lines)

    if not b64_text:
        return ConversionResult(ok=False, error="no base64 data found")

    if progress_cb:
        progress_cb(30, 100)

    try:
        raw = base64.b64decode(b64_text)
    except Exception as e:
        return ConversionResult(ok=False, error=f"decode failed: {e}")

    # 若输出路径没有扩展名，追加原始扩展名
    out_path = output_path
    base_out, out_ext = os.path.splitext(output_path)
    if not out_ext:
        out_path = output_path + src_ext

    _ensure_parent_dir(out_path)
    try:
        with open(out_path, "wb") as f:
            f.write(raw)
    except Exception as e:
        return ConversionResult(ok=False, error=f"write failed: {e}")

    if progress_cb:
        progress_cb(100, 100)

    return ConversionResult(
        ok=True, output_path=out_path,
        message=f"base64->file success: {os.path.basename(out_path)} ({len(raw)} bytes)",
        metadata={"decoded_bytes": len(raw), "original_ext": src_ext},
    )


# ============================================================
# 视频互转 (FFmpeg)
# ============================================================

# 视频容器格式 → ffmpeg 输出参数
_VIDEO_FORMAT_PARAMS = {
    "mp4": ["-c:v", "libx264", "-preset", "fast", "-c:a", "aac", "-b:a", "128k"],
    "mkv": ["-c:v", "libx264", "-preset", "fast", "-c:a", "aac", "-b:a", "128k"],
    "avi": ["-c:v", "mpeg4", "-preset", "fast", "-c:a", "mp3", "-b:a", "128k"],
    "mov": ["-c:v", "libx264", "-preset", "fast", "-c:a", "aac", "-b:a", "128k"],
    "webm": ["-c:v", "libvpx-vp9", "-preset", "fast", "-c:a", "opus"],
    "flv": ["-c:v", "flv", "-c:a", "mp3"],
    "wmv": ["-c:v", "wmv2", "-c:a", "wmav2"],
}

# 输入格式 → 容器推断（用于扩展名补全）
_INPUT_CONTAINER_MAP = {
    ".mp4": "mp4", ".m4v": "mp4",
    ".mkv": "mkv",
    ".avi": "avi",
    ".wmv": "wmv",
    ".flv": "flv",
    ".mov": "mov",
    ".webm": "webm",
    ".ts": "mp4",
    ".3gp": "mp4",
    ".m4a": "mp4",
    ".mp3": "mp4",
    ".wav": "mp4",
    ".flac": "mp4",
    ".aac": "mp4",
    ".ogg": "webm",
    ".opus": "webm",
}


def _probe_duration_ffmpeg(path: str) -> float:
    """用 ffprobe 探测媒体时长（秒），失败返回 0"""
    ffprobe_exe = shutil.which("ffprobe")
    if not ffprobe_exe:
        return 0.0
    try:
        r = subprocess.run(
            [ffprobe_exe, "-v", "quiet", "-show_entries", "format=duration",
             "-of", "default=noprint_wrappers=1:nokey=1", path],
            capture_output=True, text=True, timeout=5.0,
            startupinfo=_get_startupinfo(),
            encoding="utf-8", errors="replace",
        )
        if r.returncode == 0 and r.stdout:
            return float(r.stdout.strip().splitlines()[0])
    except Exception:
        pass
    return 0.0


def _get_startupinfo():
    """返回 subprocess.STARTUPINFO (Windows 下隐藏控制台窗口)"""
    if sys.platform == "win32":
        si = subprocess.STARTUPINFO()
        si.dwFlags |= subprocess.STARTF_USESHOWWINDOW
        return si
    return None


def convert_video_ffmpeg(
    source_path: str,
    output_path: str,
    target_format: str = "",
    progress_cb: Optional[Callable[[int, int], None]] = None,
) -> ConversionResult:
    """通过 ffmpeg 执行视频/音频格式互转

    参数：
        source_path: 源文件路径
        output_path: 输出文件路径
        target_format: 目标容器格式 (mp4/mkv/avi/mov/webm/flv/wmv)，
                       为空时从 output_path 扩展名推断
        progress_cb: 进度回调 (cur 0~100, total 100)

    返回：
        ConversionResult
    """
    if not os.path.isfile(source_path):
        return ConversionResult(ok=False, error=f"source not found: {source_path}")

    ffmpeg_exe = shutil.which("ffmpeg")
    if not ffmpeg_exe:
        return ConversionResult(ok=False, error="ffmpeg not found in PATH")

    # 推断目标格式
    if not target_format:
        ext = os.path.splitext(output_path)[1].lstrip(".").lower()
        target_format = ext
    else:
        target_format = target_format.lower()

    if target_format not in _VIDEO_FORMAT_PARAMS:
        return ConversionResult(
            ok=False,
            error=f"unsupported target format: {target_format}",
            metadata={"supported": list(_VIDEO_FORMAT_PARAMS.keys())},
        )

    # 推断输出扩展名
    _, ext = os.path.splitext(output_path)
    if not ext:
        output_path = output_path + "." + target_format
    _ensure_parent_dir(output_path)

    # 探一下源文件时长
    total_seconds = _probe_duration_ffmpeg(source_path)

    # 构造命令
    cmd = [
        ffmpeg_exe,
        "-y",
        "-i", source_path,
    ] + _VIDEO_FORMAT_PARAMS[target_format] + [output_path]

    try:
        proc = subprocess.Popen(
            cmd,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            startupinfo=_get_startupinfo(),
            encoding="utf-8",
            errors="replace",
        )
    except Exception as e:
        return ConversionResult(ok=False, error=f"ffmpeg launch failed: {e}")

    # 解析进度
    import re as _re
    _prog_re = _re.compile(r"time=(\d{2}):(\d{2}):(\d{2})\.(\d+)")

    def _reader():
        try:
            for line in proc.stderr or []:
                m = _prog_re.search(line)
                if m and total_seconds > 0:
                    h, mi, se = int(m.group(1)), int(m.group(2)), int(m.group(3))
                    ms = int(m.group(4)[:3]) if len(m.group(4)) >= 3 else 0
                    cur = h * 3600 + mi * 60 + se + ms / 1000
                    ratio = max(0.0, min(1.0, cur / total_seconds))
                    if progress_cb:
                        progress_cb(int(ratio * 100), 100)
        except Exception:
            pass

    import threading
    t = threading.Thread(target=_reader, daemon=True)
    t.start()
    proc.wait()
    t.join(timeout=5.0)

    if proc.returncode != 0:
        stderr_content = (proc.stderr or "").read() if hasattr(proc.stderr, "read") else (proc.stderr or "")
        return ConversionResult(ok=False, error=f"ffmpeg failed: {stderr_content[:300]}")

    ok = os.path.isfile(output_path) and os.path.getsize(output_path) > 0
    if not ok:
        return ConversionResult(ok=False, error="output not generated")

    return ConversionResult(
        ok=True,
        output_path=output_path,
        message=f"video->{target_format} success: {os.path.basename(output_path)}",
        metadata={"format": target_format, "output_size": os.path.getsize(output_path)},
    )


# ============================================================
# 统一派发 + 格式探测辅助
# ============================================================

# 支持的转换表 (src_fmt, dst_fmt) -> fn
SUPPORTED_CONVERSIONS = {
    # 文档类
    ("md", "html"): md_to_html,
    ("md", "docx"): md_to_docx,
    ("html", "md"): html_to_md,
    ("html", "pdf"): html_to_pdf,
    ("docx", "html"): docx_to_html,
    ("docx", "md"): docx_to_md,
    ("pdf", "txt"): pdf_to_text,
    ("pdf", "md"): pdf_to_md,
    ("pdf", "png"): pdf_to_png,
    ("pdf", "png_folder"): pdf_to_png_folder,
    # 表格类
    ("xlsx", "csv"): xlsx_to_csv,
    ("xlsx", "tsv"): xlsx_to_tsv,
    ("csv", "xlsx"): csv_to_xlsx,
    ("csv", "csv"): csv_to_csv,
    # Base64 编解码
    ("b64enc", "txt"): file_to_base64,
    ("b64dec", "bin"): base64_to_file,
}


def _detect_encoding_simple(path: str) -> str:
    """简易编码探测"""
    try:
        with open(path, "rb") as f:
            head = f.read(min(65536, os.path.getsize(path) or 0))
        head.decode("utf-8")
        return "utf-8"
    except UnicodeDecodeError:
        try:
            head.decode("gb18030")
            return "gb18030"
        except UnicodeDecodeError:
            return "utf-8"


def _detect_delimiter_simple(path: str, encoding: str) -> str:
    """简易分隔符探测"""
    try:
        with open(path, "r", encoding=encoding, errors="replace") as f:
            sample = f.read(8192)
    except OSError:
        return ","
    lines = sample.splitlines()[:5]
    if not lines:
        return ","
    for delim in ("\t", ",", ";", "|"):
        counts = [line.count(delim) for line in lines]
        if counts and all(c > 0 for c in counts) and len(set(counts)) <= 2:
            return delim
    return ","


def _is_image_format(fmt: str) -> bool:
    return fmt.lower() in IMAGE_FORMAT_MAP


def _is_base64_enc(src_fmt: str, dst_fmt: str) -> bool:
    """判断是否 Base64 编码模式（任意格式 -> txt 且源是 b64enc）"""
    return src_fmt == "b64enc"


def _is_base64_dec(src_fmt: str) -> bool:
    """判断是否 Base64 解码模式"""
    return src_fmt == "b64dec"


def convert(source_path: str, output_path: str,
            src_format: str = "", dst_format: str = "",
            progress_cb: Optional[Callable[[int, int], None]] = None) -> ConversionResult:
    """统一转换接口

    特殊约定：
      - src_format="b64enc": 文件 -> Base64 编码 TXT
      - src_format="b64dec": Base64 TXT -> 原格式文件
      - 图像互转：src_format 和 dst_format 均为图像扩展名（png, jpg 等）
    """
    # 自动推断格式
    if not src_format:
        if source_path.lower().endswith((".b64.txt", ".base64.txt")):
            src_format = "b64dec"
        else:
            src_format = os.path.splitext(source_path)[1].lstrip(".").lower()
    if not dst_format:
        dst_format = os.path.splitext(output_path)[1].lstrip(".").lower()

    # Base64 编码模式
    if src_format == "b64enc" or (dst_format == "txt" and src_format not in SUPPORTED_CONVERSIONS):
        # 通用文件 -> Base64
        return file_to_base64(source_path, output_path, progress_cb)

    # Base64 解码模式
    if src_format == "b64dec":
        return base64_to_file(source_path, output_path, progress_cb)

    # PDF 多页导出 PNG 文件夹（output_path 实际是文件夹路径）
    if src_format == "pdf" and dst_format == "png_folder":
        return pdf_to_png_folder(source_path, output_path, progress_cb=progress_cb)

    # 图像互转
    if _is_image_format(src_format) and _is_image_format(dst_format):
        return convert_image(source_path, output_path, dst_format, progress_cb=progress_cb)

    # 视频互转（源是视频/音频格式，目标是容器格式）
    _video_exts = set(k.lstrip(".") for k in _INPUT_CONTAINER_MAP.keys())
    if src_format in _video_exts and dst_format in _VIDEO_FORMAT_PARAMS:
        return convert_video_ffmpeg(source_path, output_path, dst_format, progress_cb)

    # 文档/表格类
    key = (src_format, dst_format)
    handler = SUPPORTED_CONVERSIONS.get(key)
    if handler is None:
        return ConversionResult(
            ok=False,
            error=f"unsupported conversion: {src_format}->{dst_format}",
            metadata={"supported": [f"{k[0]}->{k[1]}" for k in SUPPORTED_CONVERSIONS]},
        )

    # 有关进度回调的签名适配
    import inspect
    sig = inspect.signature(handler)
    if "progress_cb" in sig.parameters:
        return handler(source_path, output_path, progress_cb)
    return handler(source_path, output_path)


def get_supported_conversions():
    """返回支持的转换列表"""
    result = [{"from": k[0], "to": k[1]} for k in SUPPORTED_CONVERSIONS]
    # 加上 PDF -> PNG
    result.append({"from": "pdf", "to": "png", "group": "pdf"})
    result.append({"from": "pdf", "to": "png_folder", "group": "pdf"})
    # 加上图像互转（列举）
    img_exts = ["png", "jpg", "webp", "bmp", "tiff", "gif"]
    for s in img_exts:
        for d in img_exts:
            if s != d:
                result.append({"from": s, "to": d, "group": "image"})
    # 加上 base64
    result.append({"from": "*", "to": "txt", "group": "base64"})
    result.append({"from": "txt.b64", "to": "*", "group": "base64"})
    # 加上视频互转
    video_src_exts = list(_INPUT_CONTAINER_MAP.keys())
    video_dst_fmts = list(_VIDEO_FORMAT_PARAMS.keys())
    for s in video_src_exts:
        for d in video_dst_fmts:
            result.append({"from": s.lstrip("."), "to": d, "group": "video"})
    return result
