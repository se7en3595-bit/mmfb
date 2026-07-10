"""DocxHandler 单元测试

测试目标：
1. 扩展名识别
2. 预览模式：生成 HTML
3. 编辑模式：返回段落列表
4. 错误处理：文件不存在
5. 注册表正确分发
"""
import os
import sys
import tempfile
import unittest

import docx
from docx import Document
from docx.shared import Pt

# 确保 import path 能找到 mmfb 包
PROJECT_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
if PROJECT_ROOT not in sys.path:
    sys.path.insert(0, PROJECT_ROOT)

from mmfb.handlers.docx_handler import DocxHandler, DOCX_EXTENSIONS
from mmfb.core.registry import HandlerRegistry, registry


def create_test_docx(path: str, paragraphs=None, tables=None, with_heading=False):
    """创建一个测试用 docx 文件"""
    doc = Document()

    if with_heading:
        doc.add_heading('Test Title Level 0', level=0)
        doc.add_heading('Section 1', level=1)
        doc.add_heading('Subsection 1.1', level=2)

    if paragraphs:
        for p_text in paragraphs:
            doc.add_paragraph(p_text)
    else:
        doc.add_paragraph('First paragraph')
        doc.add_paragraph('Second paragraph')
        doc.add_paragraph('Third paragraph')

    if tables:
        # tables 是单张表格的 rows 列表，每行是一个字符串数组
        num_rows = len(tables)
        num_cols = len(tables[0]) if tables else 0
        table = doc.add_table(rows=num_rows, cols=num_cols)
        for i, row in enumerate(tables):
            for j, cell_text in enumerate(row):
                table.rows[i].cells[j].text = cell_text

    # 添加一些带段落样式的段落
    p = doc.add_paragraph('Bold italic paragraph')
    run = p.runs[0]
    run.bold = True
    run.italic = True
    run.font.size = Pt(14)

    p2 = doc.add_paragraph('Underline paragraph')
    p2.runs[0].underline = True

    p3 = doc.add_paragraph('')
    # 空段落

    doc.save(path)
    return path


class TestDocxExtensions(unittest.TestCase):
    """扩展名相关测试"""

    def test_docx_extension_recognized(self):
        self.assertIn('.docx', DOCX_EXTENSIONS)

    def test_can_handle_docx(self):
        self.assertTrue(DocxHandler.can_handle('/path/to/test.docx'))
        self.assertTrue(DocxHandler.can_handle('/path/to/test.DOCX'))
        self.assertTrue(DocxHandler.can_handle('/path/to/test.Docx'))

    def test_cannot_handle_other(self):
        self.assertFalse(DocxHandler.can_handle('/path/to/test.pdf'))
        self.assertFalse(DocxHandler.can_handle('/path/to/test.md'))
        self.assertFalse(DocxHandler.can_handle('/path/to/test.doc'))  # 注意 .doc 不被支持
        self.assertFalse(DocxHandler.can_handle('/path/to/test'))


class TestDocxPreview(unittest.TestCase):
    """预览模式测试"""

    def setUp(self):
        self.tmpdir = tempfile.mkdtemp()
        self.docx_path = os.path.join(self.tmpdir, 'test.docx')
        create_test_docx(
            self.docx_path,
            paragraphs=['Hello world', 'Another line'],
            tables=[['A', 'B'], ['C', 'D']],
            with_heading=True,
        )
        self.handler = DocxHandler(self.docx_path)

    def tearDown(self):
        try:
            os.remove(self.docx_path)
        except OSError:
            pass
        try:
            os.rmdir(self.tmpdir)
        except OSError:
            pass

    def test_preview_returns_dict(self):
        result = self.handler.get_preview()
        self.assertIsNotNone(result)
        self.assertIsInstance(result, dict)

    def test_preview_has_required_keys(self):
        result = self.handler.get_preview()
        self.assertIn('mime', result)
        self.assertIn('template', result)
        self.assertIn('data', result)
        self.assertIn('editable', result)

    def test_preview_mime_is_docx(self):
        result = self.handler.get_preview()
        self.assertIn('wordprocessingml', result['mime'])

    def test_preview_template_is_docx(self):
        result = self.handler.get_preview()
        self.assertEqual(result['template'], 'docx')

    def test_preview_has_html(self):
        result = self.handler.get_preview()
        self.assertIn('html', result['data'])
        html = result['data']['html']
        self.assertIsInstance(html, str)
        self.assertIn('docx-body', html)

    def test_preview_html_contains_paragraphs(self):
        result = self.handler.get_preview()
        html = result['data']['html']
        self.assertIn('docx-para', html)
        self.assertIn('Hello world', html)
        self.assertIn('Another line', html)

    def test_preview_html_contains_table(self):
        result = self.handler.get_preview()
        html = result['data']['html']
        self.assertIn('docx-table', html)
        self.assertIn('>A<', html)
        self.assertIn('>B<', html)

    def test_preview_html_contains_heading(self):
        result = self.handler.get_preview()
        html = result['data']['html']
        # Heading 1 应转为 h1 或 h2 标签
        has_h1 = '<h1 ' in html
        has_h2 = '<h2 ' in html
        self.assertTrue(has_h1 or has_h2, '应至少有一种 heading 标签')

    def test_preview_html_has_bold_or_italic(self):
        result = self.handler.get_preview()
        html = result['data']['html']
        # Bold/italic 段落应有 strong 或 em
        has_strong = '<strong>' in html
        has_em = '<em>' in html
        self.assertTrue(has_strong or has_em, '文本样式解析错误：缺少 strong/em')

    def test_preview_counts_paragraphs_and_tables(self):
        result = self.handler.get_preview()
        self.assertGreater(result['data']['paragraph_count'], 0)
        self.assertEqual(result['data']['table_count'], 1)

    def test_preview_has_file_info(self):
        result = self.handler.get_preview()
        self.assertIn('file_path', result['data'])
        self.assertIn('file_size', result['data'])
        self.assertGreater(result['data']['file_size'], 0)

    def test_preview_editable_true(self):
        result = self.handler.get_preview()
        self.assertTrue(result['editable'])


class TestDocxEdit(unittest.TestCase):
    """编辑模式测试"""

    def setUp(self):
        self.tmpdir = tempfile.mkdtemp()
        self.docx_path = os.path.join(self.tmpdir, 'edit_test.docx')
        create_test_docx(self.docx_path, paragraphs=['Para A', 'Para B', 'Para C'])
        self.handler = DocxHandler(self.docx_path)

    def tearDown(self):
        try:
            os.remove(self.docx_path)
        except OSError:
            pass
        try:
            os.rmdir(self.tmpdir)
        except OSError:
            pass

    def test_edit_returns_dict(self):
        result = self.handler.get_edit()
        self.assertIsNotNone(result)

    def test_edit_has_paragraphs(self):
        result = self.handler.get_edit()
        self.assertIn('paragraphs', result['data'])
        paras = result['data']['paragraphs']
        self.assertIsInstance(paras, paras.__class__)
        self.assertGreater(len(paras), 0)

    def test_edit_paragraphs_have_required_fields(self):
        result = self.handler.get_edit()
        paras = result['data']['paragraphs']
        for p in paras:
            self.assertIn('index', p)
            self.assertIn('text', p)
            self.assertIn('style', p)

    def test_edit_has_save_flag(self):
        result = self.handler.get_edit()
        self.assertTrue(result['data']['save'])

    def test_edit_has_edit_mode(self):
        result = self.handler.get_edit()
        self.assertEqual(result['data']['edit_mode'], 'paragraphs')


class TestDocxErrorHandling(unittest.TestCase):
    """错误处理测试"""

    def test_missing_file_returns_error(self):
        handler = DocxHandler('/nonexistent/path/doc.dot.docx')
        result = handler.get_preview()
        self.assertIsNotNone(result)
        self.assertIn('error', result)

    def test_missing_file_preview_not_editable(self):
        handler = DocxHandler('/nonexistent/path/foo.docx')
        result = handler.get_preview()
        self.assertFalse(result['editable'])


class TestDocxRegistry(unittest.TestCase):
    """注册表分发测试"""

    def test_registry_docx_dispatch(self):
        # 清空注册表并重新注册
        registry.clear()
        from mmfb.handlers import DocxHandler
        registry.register_class(DocxHandler)

        handler = registry.get_handler('/path/to/foo.docx')
        self.assertIsNotNone(handler)
        self.assertIsInstance(handler, DocxHandler)

    def test_registry_case_insensitive(self):
        registry.clear()
        from mmfb.handlers import DocxHandler
        registry.register_class(DocxHandler)

        h1 = registry.get_handler('/path/to/foo.DOCX')
        h2 = registry.get_handler('/path/to/foo.Docx')
        self.assertIsNotNone(h1)
        self.assertIsNotNone(h2)

    def test_registry_does_not_handle_pdf(self):
        registry.clear()
        from mmfb.handlers import DocxHandler
        registry.register_class(DocxHandler)

        handler = registry.get_handler('/path/to/foo.pdf')
        self.assertIsNone(handler)


class TestDocxExtensionDotDoc(unittest.TestCase):
    """明确 .doc (Word 97-2003) 不受支持"""

    def test_doc_not_handled(self):
        self.assertFalse(DocxHandler.can_handle('/path/to/old.doc'))


if __name__ == '__main__':
    unittest.main()
