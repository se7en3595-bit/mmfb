"""共享 pytest fixtures

提供临时样本文件、重置后的注册表、以及共享的测试工具函数。
"""
import os
import sys
import tempfile
from pathlib import Path

import pytest


# 项目根目录加入 sys.path，确保 import mmfb 可用
PROJECT_ROOT = Path(__file__).resolve().parent.parent.parent
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))


# ============================================================
# 样本文件生成
# ============================================================

@pytest.fixture(scope="session")
def sample_dir():
    """创建临时目录存放所有样本文件，会话级只创建一次"""
    d = tempfile.mkdtemp(prefix="mmfb_smoke_")
    yield d
    # 会话结束后清理
    import shutil
    try:
        shutil.rmtree(d, ignore_errors=True)
    except Exception:
        pass


@pytest.fixture
def sample_pdf(sample_dir):
    """生成一个最小的合法 PDF 文件（1 页空白）"""
    path = os.path.join(sample_dir, "sample.pdf")
    # 最小合法 PDF 结构
    pdf_content = (
        b"%PDF-1.4\n"
        b"1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n"
        b"2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj\n"
        b"3 0 obj<</Type/Page/MediaBox[0 0 612 792]/Parent 2 0 R>>endobj\n"
        b"xref\n0 4\n"
        b"0000000000 65535 f \n"
        b"0000000009 00000 n \n"
        b"0000000058 00000 n \n"
        b"0000000115 00000 n \n"
        b"trailer<</Size 4/Root 1 0 R>>\n"
        b"startxref\n206\n%%EOF"
    )
    with open(path, "wb") as f:
        f.write(pdf_content)
    return path


@pytest.fixture
def sample_docx(sample_dir):
    """生成一个最小的 Word docx 文件（python-docx 构建）"""
    path = os.path.join(sample_dir, "sample.docx")
    try:
        from docx import Document
        doc = Document()
        doc.add_heading("MMFB Smoke Test", level=1)
        doc.add_paragraph("Hello from smoke test paragraph.")
        doc.add_paragraph("Second paragraph with bold text.")
        doc.save(path)
    except ImportError:
        pytest.skip("python-docx not installed")
    return path


@pytest.fixture
def sample_png(sample_dir):
    """生成一个 32x32 的 PNG 图像"""
    path = os.path.join(sample_dir, "sample.png")
    try:
        from PIL import Image
        img = Image.new("RGB", (32, 32), color=(240, 220, 180))
        img.save(path, "PNG")
    except ImportError:
        pytest.skip("Pillow not installed")
    return path


@pytest.fixture
def sample_jpg(sample_dir):
    """生成一个 32x32 的 JPEG 图像"""
    path = os.path.join(sample_dir, "sample.jpg")
    try:
        from PIL import Image
        img = Image.new("RGB", (32, 32), color=(200, 180, 140))
        img.save(path, "JPEG")
    except ImportError:
        pytest.skip("Pillow not installed")
    return path


@pytest.fixture
def sample_md(sample_dir):
    """生成一个 Markdown 文件"""
    path = os.path.join(sample_dir, "sample.md")
    with open(path, "w", encoding="utf-8") as f:
        f.write("# Hello MMFB\n\nThis is a **smoke test** markdown file.\n\n- item 1\n- item 2\n")
    return path


@pytest.fixture
def sample_html(sample_dir):
    """生成一个 HTML 文件"""
    path = os.path.join(sample_dir, "sample.html")
    with open(path, "w", encoding="utf-8") as f:
        f.write("<!DOCTYPE html><html><body><h1>Test</h1><p>Smoke test</p></body></html>")
    return path


@pytest.fixture
def sample_csv(sample_dir):
    """生成 CSV 文件"""
    path = os.path.join(sample_dir, "sample.csv")
    with open(path, "w", encoding="utf-8") as f:
        f.write("name,age,city\nAlice,30,Beijing\nBob,25,Shanghai\n")
    return path


@pytest.fixture
def sample_txt(sample_dir):
    """生成纯文本文件"""
    path = os.path.join(sample_dir, "sample.txt")
    with open(path, "w", encoding="utf-8") as f:
        f.write("Hello MMFB\nThis is a smoke test file.\n")
    return path


# ============================================================
# 注册表隔离
# ============================================================

@pytest.fixture
def isolated_registry():
    """提供一个清空过的 HandlerRegistry 实例（不与全局注册表耦合）"""
    from mmfb.core.handler_base import BaseHandler
    from mmfb.core.registry import HandlerRegistry

    reg = HandlerRegistry()
    # 注册 pdf/docx/image 三个常用 handler 用于测试
    # 延迟导入，避免在 fixture 收集阶段因依赖缺失导致整个测试文件崩溃
    try:
        from mmfb.handlers.pdf_handler import PdfHandler
        reg.register_class(PdfHandler)
    except Exception:
        pass
    try:
        from mmfb.handlers.docx_handler import DocxHandler
        reg.register_class(DocxHandler)
    except Exception:
        pass
    try:
        from mmfb.handlers.image_handler import ImageHandler
        reg.register_class(ImageHandler)
    except Exception:
        pass
    try:
        from mmfb.handlers.markdown_handler import MarkdownHandler
        reg.register_class(MarkdownHandler)
    except Exception:
        pass
    try:
        from mmfb.handlers.csv_handler import CsvHandler
        reg.register_class(CsvHandler)
    except Exception:
        pass
    try:
        from mmfb.handlers.text_handler import TextHandler
        reg.register_class(TextHandler)
    except Exception:
        pass

    return reg


# ============================================================
# 测试辅助
# ============================================================

@pytest.fixture
def tmp_output(sample_dir):
    """提供一个临时输出文件路径生成器"""
    def _gen(name):
        return os.path.join(sample_dir, name)
    return _gen
