"""Word .docx 格式处理器

职责：
1. 解析 .docx 文件（段落/表格/图片/样式），输出预览数据
2. 预览模式：前端用已转换好的 HTML 渲染（只读）
3. 编辑模式：返回段落文本（textarea），保存时按"纯文本模式"
   覆盖原段落内容（保留段落级别样式）

技术方案：
- 用 python-docx 解析 docx 内部结构
- 将段落/表格/图片转换为结构化数据
- 图片以 base64 内嵌在 HTML 中
- 前端拿到 HTML 字符串直接渲染

编辑模式说明：
    第一阶段仅支持段落文本内容的编辑，保留段落/字符样式；
    表格/图片/复杂对象仅展示，不参与编辑。
    后续任务可引入 mammoth (docx.js) 作为 WYSIWYG 兜底。
"""
import base64
import json
import os
from typing import Any, Dict, List, Optional

from docx import Document
from docx.document import Document as DocumentType
from docx.table import Table
from docx.text.paragraph import Paragraph

from mmfb.core.handler_base import BaseHandler
from mmfb.core.file_handler import safe_read_text


DOCX_EXTENSIONS: List[str] = [
    ".docx",
]


class DocxHandler(BaseHandler):
    """Word .docx 文件处理器

    支持的扩展名：.docx

    预览输出结构：
    - mime: 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    - template: 'docx'
    - data.html: 完整可直接渲染的 HTML 字符串
    - data.images: [{ id, mime, base64 }]（已内嵌到 HTML，此字段备用）
    - data.paragraph_count: 段落数
    - data.table_count: 表格数
    - data.file_path: 文件路径
    - data.file_size: 字节数
    - editable: True（段落文本可编辑）
    - error: 可选错误信息
    """

    extensions = DOCX_EXTENSIONS

    def get_preview(self) -> Optional[Dict[str, Any]]:
        """获取 docx 预览数据"""
        try:
            if not os.path.isfile(self.path):
                return self._make_error_result("file not found")

            file_size = os.path.getsize(self.path)

            try:
                doc = Document(self.path)
            except Exception as e:
                return self._make_error_result(f"failed to open docx: {e}", file_size=file_size)

            # 先构建 embed 映射，供段落内联图片使用
            embeds = self._build_embed_map(doc)

            html_fragments: List[str] = []
            paragraph_count = 0
            table_count = 0

            for element in doc.element.body:
                tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag

                if tag == 'p':
                    paragraph_count += 1
                    para = Paragraph(element, doc)
                    html_fragments.append(self._paragraph_to_html(para, embeds))

                elif tag == 'tbl':
                    table_count += 1
                    table = Table(element, doc)
                    html_fragments.append(self._table_to_html(table))

            # 兼容旧前端：也返回独立的 images 数组
            images = [{"id": k, "mime": v["mime"], "base64": v["base64"]} for k, v in embeds.items()]

            # 合并 body 内部 HTML，加上 docx-viewer 包装
            body_html = '\n'.join(html_fragments)
            full_html = (
                '<div class="docx-body">'
                + body_html
                + '</div>'
            )

            return {
                "mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "template": "docx",
                "data": {
                    "html": full_html,
                    "images": images,
                    "paragraph_count": paragraph_count,
                    "table_count": table_count,
                    "file_path": os.path.abspath(self.path),
                    "file_size": file_size,
                },
                "editable": True,
            }
        except Exception as e:
            return self._make_error_result(str(e))

    def get_edit(self) -> Optional[Dict[str, Any]]:
        """获取编辑数据（段落纯文本 + 位置索引）

        编辑模式返回结构化段落数据，前端用 textarea 或多段
        输入框让用户编辑。保存时按索引重建段落文本。
        """
        preview = self.get_preview()
        if preview is None or "error" in preview:
            return preview

        try:
            doc = Document(self.path)
        except Exception:
            return preview

        paragraphs: List[Dict[str, Any]] = []
        for i, para in enumerate(doc.paragraphs):
            text = para.text or ""
            style_name = para.style.name if para.style else "Normal"
            paragraphs.append({
                "index": i,
                "text": text,
                "style": style_name,
            })

        preview["data"]["paragraphs"] = paragraphs
        preview["data"]["save"] = True
        preview["data"]["edit_mode"] = "paragraphs"
        return preview

    def save_paragraphs(self, changes_json: str) -> bool:
        """按索引回写段落文本到原 .docx 文件
        Args:
            changes_json: JSON 数组 [{index, text, style}, ...]
        Returns:
            True 成功 / False 失败
        """
        try:
            changes = json.loads(changes_json)
            if not isinstance(changes, list):
                return False
            doc = Document(self.path)
            doc_paragraphs = list(doc.paragraphs)
            for ch in changes:
                idx = ch.get("index")
                new_text = ch.get("text", "")
                if idx is None:
                    continue
                if idx < 0 or idx >= len(doc_paragraphs):
                    continue
                para = doc_paragraphs[idx]
                # 保留第一个 run 的样式，清空其余 run
                if para.runs:
                    first_run = para.runs[0]
                    for run in para.runs[1:]:
                        run.text = ""
                    first_run.text = new_text
                else:
                    # 无 run 时添加新 run
                    para.add_run(new_text)
            doc.save(self.path)
            return True
        except Exception as e:
            print(f"[DocxHandler] save_paragraphs error: {e}")
            return False

    def _paragraph_to_html(self, para: Paragraph, embeds: Optional[Dict] = None) -> str:
        """将一个 Paragraph 对象转为 HTML，内嵌 <img> 引用 data URI"""
        text = para.text or ""
        if not text.strip():
            return '<p class="docx-para docx-para--empty">&nbsp;</p>'

        style_name = (para.style.name if para.style else "Normal").lower()

        if style_name.startswith("heading 1") or style_name == "title":
            tag = "h1"
        elif style_name.startswith("heading 2"):
            tag = "h2"
        elif style_name.startswith("heading 3"):
            tag = "h3"
        elif style_name.startswith("heading 4"):
            tag = "h4"
        elif style_name.startswith("heading 5"):
            tag = "h5"
        elif style_name.startswith("heading 6"):
            tag = "h6"
        elif style_name.startswith("list paragraph") or "bullet" in style_name:
            tag = "li"
        else:
            tag = "p"

        inline_html = self._runs_to_inline_html(para, embeds)

        if not inline_html:
            inline_html = self._escape_html(text)

        css_class = "docx-para"
        if tag == "li":
            css_class += " docx-list-item"

        return f'<{tag} class="{css_class}">{inline_html}</{tag}>'

    def _runs_to_inline_html(self, para: Paragraph, embeds: Optional[Dict] = None) -> str:
        """将 Paragraph 内部的 runs 转为带样式的 HTML，内嵌图片转为 <img>"""
        if not para.runs:
            return ""

        embeds = embeds or {}
        nsmap = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
            'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
        }

        fragments: List[str] = []
        for run in para.runs:
            r_elem = run._element
            # 检查 run 内是否有内联图片 (w:drawing)
            drawings = r_elem.findall('.//w:drawing', nsmap)
            if drawings:
                for drawing in drawings:
                    blips = drawing.findall('.//a:blip', nsmap)
                    for blip in blips:
                        embed_id = blip.get(
                            '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                        if embed_id and embed_id in embeds:
                            img = embeds[embed_id]
                            src = f"data:{img['mime']};base64,{img['base64']}"
                            fragments.append(
                                f'<img src="{src}" class="docx-img" style="max-width:100%;" />')
            else:
                text = run.text or ""
                if not text:
                    continue
                html = self._escape_html(text)

                bold = run.bold
                italic = run.italic
                underline = run.underline
                strikethrough = run.font.strike if run.font else False

                if bold:
                    html = f"<strong>{html}</strong>"
                if italic:
                    html = f"<em>{html}</em>"
                if underline:
                    html = f'<span style="text-decoration:underline">{html}</span>'
                if strikethrough:
                    html = f"<del>{html}</del>"

                fragments.append(html)

        return "".join(fragments)

    def _table_to_html(self, table: Table) -> str:
        """将一个 Table 对象转为 HTML table"""
        rows_html: List[str] = []
        for row in table.rows:
            cells_html: List[str] = []
            for cell in row.cells:
                cell_text = (cell.text or "").strip()
                escaped = self._escape_html(cell_text) or "&nbsp;"
                cells_html.append(f"<td>{escaped}</td>")
            rows_html.append("<tr>" + "".join(cells_html) + "</tr>")
        return (
            '<table class="docx-table">'
            + "\n".join(rows_html)
            + '</table>'
        )

    def _build_embed_map(self, doc: DocumentType) -> Dict[str, Dict[str, str]]:
        """构建 rEmbedID -> {mime, base64} 映射，用于段落内联图片"""
        embed_map: Dict[str, Dict[str, str]] = {}
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                try:
                    blob = rel.target_part.blob
                    mime = rel.target_part.content_type or "image/png"
                    b64 = base64.b64encode(blob).decode("ascii")
                    embed_map[rel.rId] = {"mime": mime, "base64": b64}
                except Exception:
                    continue
        return embed_map

    def _extract_images(self, doc: DocumentType) -> List[Dict[str, str]]:
        """从 docx 中提取所有图片，返回 [{ id, mime, base64 }]"""
        images: List[Dict[str, str]] = []
        image_idx = 0
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                try:
                    blob = rel.target_part.blob
                    mime = rel.target_part.content_type or "image/png"
                    b64 = base64.b64encode(blob).decode("ascii")
                    images.append({
                        "id": f"img_{image_idx}",
                        "mime": mime,
                        "base64": b64,
                    })
                    image_idx += 1
                except Exception:
                    continue
        return images

    def _make_error_result(self, error: str, file_size: int = 0) -> Dict[str, Any]:
        """构造错误结果"""
        return {
            "mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "template": "docx",
            "data": {
                "html": f'<div class="docx-body"><p class="docx-error">{self._escape_html(error)}</p></div>',
                "images": [],
                "paragraph_count": 0,
                "table_count": 0,
                "file_path": os.path.abspath(self.path),
                "file_size": file_size,
            },
            "editable": False,
            "error": error,
        }

    @staticmethod
    def _escape_html(text: str) -> str:
        """HTML 实体转义"""
        return (
            str(text)
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;")
            .replace("'", "&#39;")
        )
